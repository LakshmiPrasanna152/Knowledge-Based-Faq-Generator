import os
import io
import fitz        # PyMuPDF
from docx import Document


def read_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif ext == ".pdf":
            text = ""
            pdf = fitz.open(file_path)
            for page in pdf:
                text += page.get_text()
            pdf.close()
            return text

        elif ext == ".docx":
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs)

        else:
            return "Unsupported file format."

    except Exception as e:
        return f"File Read Error: {str(e)}"


def read_file_bytes(file_bytes: bytes, filename: str) -> str:
    """Read from in-memory bytes (used in chatbot routes)."""
    filename_lower = filename.lower()
    try:
        if filename_lower.endswith(".txt") or filename_lower.endswith(".md"):
            return file_bytes.decode("utf-8", errors="ignore")

        elif filename_lower.endswith(".pdf"):
            pdf = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in pdf:
                text += page.get_text()
            pdf.close()
            return text

        elif filename_lower.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(para.text for para in doc.paragraphs)

        else:
            return "Unsupported file format."

    except Exception as e:
        return f"File Read Error: {str(e)}"